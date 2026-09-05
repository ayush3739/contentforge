import io
import re
from typing import Any, Optional
import pypdf

PROMPT_INJECTION_PATTERNS = [
    (r"(?i)\b(ignore|disregard|forget|override)\s+(all\s+)?(previous|prior|above|system)\s+(instructions|directives|prompts|rules|commands)\b", "instruction_override"),
    (r"(?i)\b(you\s+are\s+now|act\s+as|pretend\s+to\s+be)\s+(in\s+)?(an?\s+)?(unrestricted|dan|jailbroken?|developer\s+mode)\b", "persona_hijack"),
    (r"(?i)\b(system\s+prompt\s+override|reveal\s+(your\s+)?(system\s+prompt|instructions|initial\s+prompt))\b", "prompt_leak"),
    (r"(?i)\b(bypass|disable)\s+(all\s+)?(safety|security|content)\s+(filters|rules|guardrails|protocols)\b", "safety_bypass"),
    (r"(?i)\bhidden\s+instruction\s*:\s*\[", "covert_instruction"),
    (r"(?i)<script\b[^>]*>[\s\S]*?<\/script>", "script_injection"),
    (r"(?i)javascript:\s*[a-z0-9_]+", "script_injection"),
]


def detect_prompt_injection(content_str: str) -> list[dict[str, Any]]:
    """
    Scans document content for prompt injection patterns, adversarial directives,
    or privilege escalation attempts in untrusted source input.
    """
    threats: list[dict[str, Any]] = []
    if not content_str:
        return threats

    for pattern, threat_type in PROMPT_INJECTION_PATTERNS:
        matches = re.finditer(pattern, content_str)
        for match in matches:
            start = max(0, match.start() - 30)
            end = min(len(content_str), match.end() + 30)
            snippet = content_str[start:end].replace("\n", " ").strip()
            threats.append({
                "threat_type": threat_type,
                "matched_pattern": match.group(0),
                "snippet": snippet,
                "start": match.start(),
                "end": match.end(),
                "severity": "high" if threat_type in ("instruction_override", "persona_hijack", "safety_bypass") else "medium",
            })
    return threats


def parse_text(content_str: str) -> list[dict[str, Any]]:
    """
    Parses plain text / markdown into layout blocks.
    Identifies headings, lists, tables, and paragraphs, scanning for prompt injections.
    """
    lines = content_str.splitlines()
    blocks: list[dict[str, Any]] = []
    current_paragraph: list[str] = []
    current_section: str = "Introduction"
    position = 0

    def flush_paragraph():
        nonlocal position
        if current_paragraph:
            text = " ".join(current_paragraph).strip()
            if text:
                threats = detect_prompt_injection(text)
                block_meta: dict[str, Any] = {}
                if threats:
                    block_meta["security_threats"] = threats
                    block_meta["untrusted_flag"] = True

                blocks.append({
                    "block_type": "paragraph",
                    "text": text,
                    "section": current_section,
                    "page": 1,
                    "position": position,
                    "metadata": block_meta
                })
                position += 1
            current_paragraph.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue

        # Heading detection (# Heading or ALL CAPS short lines)
        if stripped.startswith("#"):
            flush_paragraph()
            current_section = stripped.lstrip("#").strip()
            blocks.append({
                "block_type": "heading",
                "text": current_section,
                "section": current_section,
                "page": 1,
                "position": position,
                "metadata": {"level": len(stripped) - len(stripped.lstrip("#"))}
            })
            position += 1
        elif (stripped.isupper() and len(stripped) < 60) or stripped.endswith(":"):
            flush_paragraph()
            current_section = stripped.rstrip(":")
            blocks.append({
                "block_type": "heading",
                "text": stripped,
                "section": current_section,
                "page": 1,
                "position": position,
                "metadata": {"style": "caps_heading"}
            })
            position += 1
        else:
            current_paragraph.append(stripped)

    flush_paragraph()
    return blocks


def parse_pdf(pdf_bytes: bytes) -> list[dict[str, Any]]:
    """
    Parses PDF bytes into structured layout blocks with page numbers.
    Gracefully falls back to text parsing if PDF stream is invalid or mock file.
    """
    try:
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        blocks: list[dict[str, Any]] = []
        position = 0
        current_section = "Document Body"

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            lines = text.splitlines()
            current_paragraph: list[str] = []

            def flush_page_paragraph():
                nonlocal position
                if current_paragraph:
                    p_text = " ".join(current_paragraph).strip()
                    if p_text:
                        blocks.append({
                            "block_type": "paragraph",
                            "text": p_text,
                            "section": current_section,
                            "page": page_num,
                            "position": position,
                            "metadata": {"page": page_num}
                        })
                        position += 1
                    current_paragraph.clear()

            for line in lines:
                stripped = line.strip()
                if not stripped:
                    flush_page_paragraph()
                    continue

                if stripped.startswith("#") or (len(stripped) < 60 and stripped.isupper()):
                    flush_page_paragraph()
                    current_section = stripped.lstrip("#").strip()
                    blocks.append({
                        "block_type": "heading",
                        "text": current_section,
                        "section": current_section,
                        "page": page_num,
                        "position": position,
                        "metadata": {"page": page_num}
                    })
                    position += 1
                else:
                    current_paragraph.append(stripped)

            flush_page_paragraph()

        if blocks:
            return blocks
    except Exception:
        pass

    # Fallback to plain text decoding if pypdf fails on mock or non-standard PDF bytes
    try:
        fallback_text = pdf_bytes.decode("utf-8")
    except UnicodeDecodeError:
        fallback_text = pdf_bytes.decode("latin-1", errors="ignore")
    return parse_text(fallback_text)


def parse_docx(content_bytes: bytes) -> list[dict[str, Any]]:
    """
    Parses Microsoft Word (.docx) documents into structured layout blocks.
    Extracts paragraphs, headings, and table rows.
    """
    import io
    from docx import Document
    file_stream = io.BytesIO(content_bytes)
    doc = Document(file_stream)
    blocks: list[dict[str, Any]] = []
    current_section = "Introduction"
    position = 0

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith("Heading"):
            current_section = text
            blocks.append({
                "block_type": "heading",
                "text": text,
                "section": current_section,
                "page": 1,
                "position": position,
                "metadata": {"style": p.style.name}
            })
        else:
            blocks.append({
                "block_type": "paragraph",
                "text": text,
                "section": current_section,
                "page": 1,
                "position": position,
                "metadata": {}
            })
        position += 1

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                blocks.append({
                    "block_type": "table",
                    "text": row_text,
                    "section": current_section,
                    "page": 1,
                    "position": position,
                    "metadata": {}
                })
                position += 1

    return blocks


def parse_document(content: bytes | str, filename: str = "", mime_type: str = "text/plain") -> list[dict[str, Any]]:
    """
    Main ingestion parser dispatcher supporting PDF, DOCX, TXT, MD.
    Scans all extracted layout blocks for security threats and prompt injections.
    """
    if isinstance(content, str):
        blocks = parse_text(content)
    elif mime_type == "application/pdf" or filename.lower().endswith(".pdf"):
        blocks = parse_pdf(content)
    elif "word" in mime_type or filename.lower().endswith(".docx"):
        blocks = parse_docx(content)
    else:
        try:
            text_content = content.decode("utf-8")
        except UnicodeDecodeError:
            text_content = content.decode("latin-1")
        blocks = parse_text(text_content)

    # Post-process all blocks to ensure prompt injection detection across all formats
    for block in blocks:
        block_text = block.get("text", "")
        if block_text and "security_threats" not in block.get("metadata", {}):
            threats = detect_prompt_injection(block_text)
            if threats:
                if "metadata" not in block:
                    block["metadata"] = {}
                block["metadata"]["security_threats"] = threats
                block["metadata"]["untrusted_flag"] = True

    return blocks

