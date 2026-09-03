"""
Sample Document Generator for ContentForge AI (SIH26154)

Generates realistic, multi-page test documents (.docx and .pdf)
modeling an enterprise cyber incident response and infrastructure report.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(BASE_DIR, "sample_documents")
os.makedirs(DOCS_DIR, exist_ok=True)


def generate_docx():
    docx_path = os.path.join(DOCS_DIR, "cyber_incident_report_INC88412.docx")
    doc = Document()

    title = doc.add_paragraph()
    title_run = title.add_run("NATIONAL CRITICAL INFRASTRUCTURE PROTECTION AGENCY (NCIPA)\n")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    sub = title.add_run("Incident Investigation Report: Operation BlackEcho (INC-88412)\nTLP:AMBER | Public Sector Resilience Assessment")
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph("―" * 55)

    doc.add_heading("1. Executive Summary & Incident Timeline", level=1)
    doc.add_paragraph(
        "On August 14, 2026, at approximately 02:41 UTC, the Central Operations Security Center (COSC) "
        "detected anomalous outbound traffic originating from core telemetry gateway servers in Cluster Delta-7. "
        "Subsequent forensic analysis confirmed unauthorized access by state-nexus threat group APT-39 (SilentHydra). "
        "The attack led to the compromise of 14 core supervisory control nodes and the exfiltration of 450 GB of "
        "operational telemetry archives."
    )

    doc.add_heading("Key Incident Metrics", level=2)
    metrics_table = doc.add_table(rows=6, cols=2)
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("Incident Identifier", "INC-88412 (Operation BlackEcho)"),
        ("Date of Initial Incursion", "August 14, 2026 (02:41 UTC)"),
        ("Compromised Infrastructure", "14 Telemetry & SCADA Gateway Servers"),
        ("Data Exfiltrated", "450 Gigabytes of telemetry archives"),
        ("Service Interruption", "4 hours 18 minutes partial telemetry blackout"),
        ("Total Estimated Financial Exposure", "$2,850,000 USD (Remediation & Forensic Audits)"),
    ]
    for i, (k, v) in enumerate(data):
        metrics_table.rows[i].cells[0].text = k
        metrics_table.rows[i].cells[1].text = v
        metrics_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    doc.add_heading("2. Technical Root Cause & Attack Vector Analysis", level=1)
    doc.add_paragraph(
        "The initial entry point was traced to an external-facing edge bastion host running a vulnerable "
        "utility layer. Specifically, the adversaries exploited CVE-2024-3094 (liblzma/XZ Utils backdoor) "
        "to establish reverse SSH tunnels bypass perimeter firewalls."
    )
    doc.add_paragraph(
        "Following initial persistence, the threat actors leveraged CVE-2024-21413 (Microsoft Outlook NTLM Credential "
        "Leakage) to harvest service account credentials belonging to svc_telemetry_admin. Privilege escalation was achieved "
        "within 38 minutes of landing on the internal subnet."
    )

    doc.add_heading("Indicators of Compromise (IoCs)", level=2)
    ioc_table = doc.add_table(rows=5, cols=3)
    ioc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ioc_headers = ["Indicator Type", "Value / Artifact", "Associated Threat Action"]
    for j, h in enumerate(ioc_headers):
        ioc_table.rows[0].cells[j].text = h
        ioc_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

    ioc_data = [
        ("IPv4 Address", "198.51.100.47", "C2 Command & Control Beacon"),
        ("SHA-256 Hash", "e3b0c44298fc1c149afbf4c8996fb92427ae41e4649b934ca495991b7852b855", "Backdoored liblzma.so.5.6.0 binary"),
        ("Mutual TLS Cert", "CN=telemetry-gateway.internal.bad", "Man-in-the-Middle proxy interception"),
        ("File Path", "/usr/local/bin/.blackecho_agent", "Obfuscated persistence beacon script"),
    ]
    for i, row in enumerate(ioc_data, start=1):
        for j, val in enumerate(row):
            ioc_table.rows[i].cells[j].text = val

    doc.add_page_break()

    doc.add_heading("3. Affected Systems & Quantitative Impact Assessment", level=1)
    doc.add_paragraph(
        "The breach impact was concentrated across the eastern regional operational division. While safety-instrumented "
        "systems (SIS) remained isolated via hardware air-gaps, operational visibility was impaired across 3 regional dispatch centers."
    )

    doc.add_heading("Impact Breakdown by Facility", level=2)
    impact_table = doc.add_table(rows=4, cols=4)
    impact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    impact_headers = ["Facility Name", "Impacted Nodes", "Outage Duration", "Restoration Cost"]
    for j, h in enumerate(impact_headers):
        impact_table.rows[0].cells[j].text = h
        impact_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

    impact_data = [
        ("Substation Apex-1", "6 SCADA Nodes", "4h 18m", "$1,120,000 USD"),
        ("Telemetry Relay East", "5 Distribution Relays", "2h 45m", "$890,000 USD"),
        ("Central Dispatch South", "3 Gateway Servers", "1h 10m", "$840,000 USD"),
    ]
    for i, row in enumerate(impact_data, start=1):
        for j, val in enumerate(row):
            impact_table.rows[i].cells[j].text = val

    doc.add_page_break()

    doc.add_heading("4. Containment Actions & Strategic Hardening Mandates", level=1)
    doc.add_paragraph(
        "The containment team executed emergency isolation protocols at 06:59 UTC on August 14, 2026. "
        "All egress routes to known command-and-control infrastructure were blocked at the national backbone BGP level."
    )
    doc.add_paragraph(
        "The following immediate remediation actions have been completed:\n"
        "1. Revoked all 84 kerberos service account tickets and enforced mandatory 24-character key rotation.\n"
        "2. Decommissioned all 14 affected server instances and rebuilt systems from verified cryptographic golden images.\n"
        "3. Patched CVE-2024-3094 across 100% of internal Linux nodes (verified 1,240 servers).\n"
        "4. Enforced hardware FIDO2 MFA across all administrative ingress points."
    )

    doc.add_heading("5. Strategic Recommendations for Leadership", level=1)
    doc.add_paragraph(
        "To ensure permanent resilience and prevent recurrence of similar advanced persistent threats, "
        "the advisory board recommends the following three core investments:\n\n"
        "• Mandate Zero Trust Micro-segmentation across OT/IT boundary layers within 60 days.\n"
        "• Establish Immutable Blockchain Provenance logging for all operational telemetry artifacts.\n"
        "• Allocate an emergency capital reserve of $1,500,000 USD for continuous automated red-team simulations."
    )

    doc.save(docx_path)
    return docx_path


def generate_pdf():
    pdf_path = os.path.join(DOCS_DIR, "cyber_incident_report_INC88412.pdf")
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=45,
        leftMargin=45,
        topMargin=45,
        bottomMargin=45,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=6,
    )
    h1_style = ParagraphStyle(
        'SectionH1',
        parent=styles['Heading2'],
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#1e3a8a'),
        spaceBefore=14,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontSize=9.5,
        leading=14,
        textColor=colors.HexColor('#334155'),
        spaceAfter=8,
    )
    meta_style = ParagraphStyle(
        'MetaStyle',
        parent=styles['Normal'],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#64748b'),
        spaceAfter=12,
    )

    story = []

    story.append(Paragraph("NATIONAL CRITICAL INFRASTRUCTURE PROTECTION AGENCY", title_style))
    story.append(Paragraph("Incident Investigation Report: Operation BlackEcho (INC-88412)<br/><b>TLP:AMBER | Public Sector Resilience Assessment</b>", meta_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph("1. Executive Summary & Incident Timeline", h1_style))
    story.append(Paragraph(
        "On August 14, 2026, at approximately 02:41 UTC, the Central Operations Security Center (COSC) "
        "detected anomalous outbound traffic originating from core telemetry gateway servers in Cluster Delta-7. "
        "Subsequent forensic analysis confirmed unauthorized access by state-nexus threat group APT-39 (SilentHydra). "
        "The attack led to the compromise of 14 core supervisory control nodes and the exfiltration of 450 GB of "
        "operational telemetry archives.",
        body_style
    ))

    metrics_data = [
        ["Incident Identifier", "INC-88412 (Operation BlackEcho)"],
        ["Date of Incursion", "August 14, 2026 (02:41 UTC)"],
        ["Compromised Infrastructure", "14 Telemetry & SCADA Gateway Servers"],
        ["Data Exfiltrated", "450 Gigabytes of telemetry archives"],
        ["Service Interruption", "4 hours 18 minutes partial telemetry blackout"],
        ["Total Financial Exposure", "$2,850,000 USD (Remediation & Forensic Audits)"],
    ]
    t1 = Table(metrics_data, colWidths=[200, 320])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#0f172a')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
    ]))
    story.append(t1)
    story.append(PageBreak())

    story.append(Paragraph("2. Technical Root Cause & Attack Vector Analysis", h1_style))
    story.append(Paragraph(
        "The initial entry point was traced to an external-facing edge bastion host running a vulnerable "
        "utility layer. Specifically, the adversaries exploited CVE-2024-3094 (liblzma/XZ Utils backdoor) "
        "to establish reverse SSH tunnels bypassing perimeter firewalls. Following initial persistence, "
        "the threat actors leveraged CVE-2024-21413 (Microsoft Outlook NTLM Credential Leakage) to harvest "
        "service account credentials belonging to svc_telemetry_admin. Privilege escalation was achieved "
        "within 38 minutes of landing on the internal subnet.",
        body_style
    ))

    ioc_data = [
        ["Indicator Type", "Value / Artifact", "Threat Action"],
        ["IPv4 Address", "198.51.100.47", "C2 Command & Control Beacon"],
        ["SHA-256 Hash", "e3b0c44298fc1c149afbf4c8996fb92427ae41e4...", "Backdoored liblzma.so.5.6.0 binary"],
        ["TLS Cert", "CN=telemetry-gateway.internal.bad", "Man-in-the-Middle proxy interception"],
        ["File Path", "/usr/local/bin/.blackecho_agent", "Obfuscated persistence beacon script"],
    ]
    t2 = Table(ioc_data, colWidths=[120, 240, 160])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a8a')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8.5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
    ]))
    story.append(t2)
    story.append(PageBreak())

    story.append(Paragraph("3. Affected Systems & Quantitative Impact Assessment", h1_style))
    story.append(Paragraph(
        "The breach impact was concentrated across the eastern regional operational division. While safety-instrumented "
        "systems (SIS) remained isolated via hardware air-gaps, operational visibility was impaired across 3 regional dispatch centers.",
        body_style
    ))

    impact_data = [
        ["Facility Name", "Impacted Nodes", "Outage Duration", "Restoration Cost"],
        ["Substation Apex-1", "6 SCADA Nodes", "4h 18m", "$1,120,000 USD"],
        ["Telemetry Relay East", "5 Distribution Relays", "2h 45m", "$890,000 USD"],
        ["Central Dispatch South", "3 Gateway Servers", "1h 10m", "$840,000 USD"],
    ]
    t3 = Table(impact_data, colWidths=[140, 120, 120, 140])
    t3.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#334155')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8.5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
    ]))
    story.append(t3)
    story.append(PageBreak())

    story.append(Paragraph("4. Containment Actions & Strategic Hardening Mandates", h1_style))
    story.append(Paragraph(
        "The containment team executed emergency isolation protocols at 06:59 UTC on August 14, 2026. "
        "All egress routes to known command-and-control infrastructure were blocked at the national backbone BGP level.<br/><br/>"
        "<b>Immediate Actions Completed:</b><br/>"
        "1. Revoked all 84 kerberos service account tickets and enforced mandatory 24-character key rotation.<br/>"
        "2. Decommissioned all 14 affected server instances and rebuilt systems from verified cryptographic golden images.<br/>"
        "3. Patched CVE-2024-3094 across 100% of internal Linux nodes (verified 1,240 servers).<br/>"
        "4. Enforced hardware FIDO2 MFA across all administrative ingress points.<br/><br/>"
        "<b>Strategic Recommendations for Leadership:</b><br/>"
        "• Mandate Zero Trust Micro-segmentation across OT/IT boundary layers within 60 days.<br/>"
        "• Establish Immutable Blockchain Provenance logging for all operational telemetry artifacts.<br/>"
        "• Allocate an emergency capital reserve of $1,500,000 USD for continuous automated red-team simulations.",
        body_style
    ))

    doc.build(story)
    return pdf_path


if __name__ == "__main__":
    generate_docx()
    generate_pdf()
